import streamlit as st
import os
from src.SDLC.ui.uiconfig import Config
from src.SDLC.LLM.configllm import LLMconfig
from src.SDLC.ui.display_result import DisplayResult
from src.SDLC.nodes.sdlc_nodes import SDLC_Nodes
from docx import Document
from dotenv import load_dotenv


def initialize_streamlit_ui(st_session_state):
        if "state" not in st_session_state:
            st_session_state.state = {"stage": "config", "data": {}}
        load_dotenv()


class LoadStreamlitUi:
    def __init__(self,st_session_state):
        self.config = Config()
        self.display = DisplayResult()
        self.llm_config = self.display.graph_builder.sdlc_nodes.llm_config
        self.stage = st_session_state.state["stage"]
        self.data = st_session_state.state["data"]

    def update_stage_and_data(self,st_session_state):
        self.stage = st_session_state.state["stage"]
        self.data = st_session_state.state["data"]

    def get_session_stage(self):
        return self.stage

    def render_streamlit_ui(self,st_session_state):
        # --- Stage 1: Setup ---
        if self.stage == "setup":
            st.header("📋 Project Setup")
            user_input = st.text_area("Enter software requirements to begin:", height=200)
            if st.button("🚀 Start AI Pipeline"):
                output = self.display.start_pipeline(user_input)
                self.data.update(output)
                st_session_state.state = {"stage": "story_review", "data": self.data}
                self.update_stage_and_data(st_session_state)
                st.rerun()

        # --- Stage 2: User Story Review ---
        elif self.stage == "story_review":
            st.header("📘 Review User Stories")
            for story in self.data["user_stories"]:
                st.markdown(f"- {story}")
            
            #---Download the stories generated------
            if self.data.get("user_stories"):
                doc = Document()
                doc.add_heading("User Stories", 0)
                for story in self.data["user_stories"]:
                    doc.add_paragraph(f"- {story}")
                doc_path = "user_stories.docx"
                doc.save(doc_path)

                with open(doc_path, "rb") as f:
                    st.download_button("📥 Download User Stories", f, file_name=doc_path)

            approval = st.radio("Approve user stories?", ["Approved", "Denied"])
            feedback = st.text_area("Optional feedback:", disabled=approval == "Approved")
            if st.button("Submit Review"):
                output = self.display.resume_pipeline({
                    "approval": approval,
                    "feedback": feedback if approval == "Denied" else ""
                })
                self.data.update(output)
                if approval == "Denied":
                    st_session_state.state = {"stage": "story_review", "data": self.data}
                else:
                    st.session_state.state = {"stage": "design_review", "data": self.data}

                self.update_stage_and_data(st_session_state)
                st.rerun()

        # --- Stage 3: Design Document Review ---
        elif self.stage == "design_review":
            st.header("📄 Review Design Document")
            st.markdown(self.data["final_doc"])

            #---Download the design document generated------
            if self.data.get("final_doc"):
                doc = Document()
                doc.add_heading("Design Document", 0)
                for section in self.data["final_doc"].split("\n\n---\n\n"):
                    doc.add_paragraph(section)
                doc_path = "design_doc.docx"
                doc.save(doc_path)

                with open(doc_path, "rb") as f:
                    st.download_button("📥 Download Design Document", f, file_name=doc_path)


            approval = st.radio("Approve design document?", ["Approved", "Denied"])
            feedback = st.text_area("Optional design feedback:", disabled=approval == "Approved")
            if st.button("Submit Design Review"):
                output = self.display.resume_pipeline({
                    "design_approval": approval,
                    "design_feedback": feedback if approval == "Denied" else ""
                })
                self.data.update(output)

                if approval == "Denied":
                    st.session_state.state = {"stage": "design_review", "data": self.data}
                else:
                    st.session_state.state = {"stage": "code_review", "data": self.data}

                self.update_stage_and_data(st_session_state)
                st.rerun()

        # --- Stage 4: Code Review ---
        elif self.stage == "code_review":
            st.header("💻 Review Code")
            st.code(self.data["source_code"], language="python")

            #---Download the source code generated------
            if self.data.get("source_code"):
                st.download_button(
                    "📥 Download Source Code", 
                    self.data["source_code"], 
                    file_name="generated_code.py"
                )

            approval = st.radio("Approve generated code?", ["Approved", "Denied"])
            feedback = st.text_area("Code feedback:", disabled=approval == "Approved")
            if st.button("Submit Code Review"):
                output = self.display.resume_pipeline({
                    "code_approval": approval,
                    "review_comments": feedback if approval == "Denied" else ""
                })
                self.data.update(output)
                st_session_state.state = {"stage": "security_review", "data": self.data}

                self.update_stage_and_data(st_session_state)
                st.rerun()

        # --- Stage 5: Security Review ---
        elif self.stage == "security_review":
            st.header("🔐 Security Review")
            st.code(self.data["source_code"], language="python")

            #---Download the source code generated------
            if self.data.get("source_code"):
                st.download_button(
                    "📥 Download Source Code", 
                    self.data["source_code"], 
                    file_name="generated_code.py"
                )
            
            approval = st.radio("Approve code from security view?", ["Approved", "Denied"])
            feedback = st.text_area("Security feedback:", disabled=approval == "Approved")
            if st.button("Submit Security Review"):
                output = self.display.resume_pipeline({
                    "code_approval": approval,
                    "review_comments": feedback if approval == "Denied" else ""
                })
                self.data.update(output)
                st_session_state.state = {"stage": "test_review", "data": self.data}

                self.update_stage_and_data(st_session_state)
                st.rerun()

        # --- Stage 6: Test Case Review ---
        elif self.stage == "test_review":
            st.header("🧪 Test Cases")
            st.code(self.data["testcases"], language="python")


            #---Download the testcase pytest file generated------
            if self.data.get("testcases"):
                st.download_button(
                    "📥 Download Test Cases", 
                    self.data["testcases"], 
                    file_name="test_cases.py"
                )

            approval = st.radio("Approve test cases?", ["Approved", "Denied"])
            feedback = st.text_area("Testcase feedback:", disabled=approval == "Approved")
            if st.button("Submit Testcase Review"):
                self.display.resume_pipeline({
                    "testcase_approval": approval,
                    "testcase_comments": feedback if approval == "Denied" else ""
                })
                st.success("🎉 All steps complete!")
                st_session_state.state = {"stage": "completed", "data": {}}

        elif self.stage == "completed":
            st.balloons()
            st.success("Process finished successfully.")
            if st.button("🔄 Restart"):
                st_session_state.state = {"stage": "setup", "data": {}}

                self.update_stage_and_data(st_session_state)
                st.rerun()

    def load_streamlit_ui(self,st_session_state):
        with st.sidebar:
            st.set_page_config(layout="wide")
            st.header("⚙️ LLM Configuration")

            st.markdown("### 🧠 Choose Models")
            model_config = {
                "user_story_model": st.selectbox("User Story Generator", self.config.get_llm_options()),
                "doc_creator_model": st.selectbox("Document Creator", self.config.get_llm_options()),
                "code_gen_model": st.selectbox("Code Generator", self.config.get_llm_options()),
                "test_writer_model": st.selectbox("Test Case Writer", self.config.get_llm_options()),
            }

            st.markdown("### 🔐 API Keys")
            openai_key = st.text_input("OpenAI API Key", type="password")
            groq_key = st.text_input("Groq API Key", type="password")
        
            st_session_state.button = False

            if st.button("💾 Save & Start"):
                st_session_state.button = True
                if openai_key == "" or groq_key =="":
                    os.environ['OPENAI_API_KEY'] = os.getenv("OPENAI_API_KEY")
                    os.environ['GROQ_API_KEY'] = os.getenv("GROQ_API_KEY")
                else:
                    os.environ["OPENAI_API_KEY"] = openai_key
                    os.environ["GROQ_API_KEY"] = groq_key

                st_session_state.model_config = model_config
                self.llm_config.configure_llms(model_config)  # <- inject into pipeline
                st_session_state.state = {"stage": "setup", "data": {}}
                self.stage = st_session_state.state["stage"]
        
        if st_session_state.button == True:
            self.render_streamlit_ui(st_session_state)

